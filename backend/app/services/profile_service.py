import asyncio
import io
import json
import uuid
from urllib.parse import urlparse
from fastapi import HTTPException, status
import fitz  # PyMuPDF
from docx import Document
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
from typing import List, Optional
from sqlalchemy.future import select
from app.models.profile import MasterProfile
from app.core.config import settings

client = genai.Client(api_key=settings.GEMINI_API_KEY)

class ExtractedBullet(BaseModel):
    id: str = Field(default_factory=lambda: f"blt_{uuid.uuid4().hex[:8]}")
    text: str

class ExtractedExperience(BaseModel):
    id: str = Field(default_factory=lambda: f"exp_{uuid.uuid4().hex[:8]}")
    company: str
    title: str
    start_date: str
    end_date: str
    bullets: List[ExtractedBullet]

class ExtractedEducation(BaseModel):
    id: str = Field(default_factory=lambda: f"edu_{uuid.uuid4().hex[:8]}")
    institution: str
    degree: str
    field_of_study: str
    start_year: str
    end_year: str

class ExtractedSkill(BaseModel):
    id: str = Field(default_factory=lambda: f"skl_{uuid.uuid4().hex[:8]}")
    name: str

class MasterProfileExtraction(BaseModel):
    full_name: str
    title: str
    email: str
    phone: str
    location: str
    bio: str
    experiences: List[ExtractedExperience]
    education: List[ExtractedEducation]
    skills: List[ExtractedSkill]

FALLBACK_MODELS = ['gemini-3.6-flash']

COMMON_SKILLS = [
    "Python", "JavaScript", "TypeScript", "Dart", "Flutter", "React", "React Native",
    "Node.js", "Express", "FastAPI", "Django", "Flask", "SQL", "PostgreSQL", "MySQL",
    "MongoDB", "Redis", "Docker", "Kubernetes", "AWS", "Google Cloud", "Azure", "Git",
    "GitHub", "CI/CD", "REST API", "GraphQL", "HTML", "CSS", "TailwindCSS", "Java",
    "Kotlin", "Swift", "C++", "C#", ".NET", "Linux", "Terraform", "Figma", "Agile", "Scrum"
]

def _call_gemini_extraction_sync(prompt: str) -> str:
    last_error = None
    for model_name in FALLBACK_MODELS:
        for attempt in range(2):
            try:
                response = client.models.generate_content(
                    model=model_name,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        response_mime_type="application/json",
                        response_schema=MasterProfileExtraction,
                        temperature=0.1,
                    ),
                )
                if response.text and response.text.strip():
                    return response.text
            except Exception as e:
                last_error = e
                err_msg = str(e)
                print(f"[Gemini Extraction] Model '{model_name}' attempt {attempt + 1} failed: {err_msg}")
                if "503" in err_msg or "high demand" in err_msg.lower() or "429" in err_msg:
                    import time
                    time.sleep(1.5)
                    continue
                else:
                    break
    if last_error:
        raise last_error
    return ""

def _extract_name_from_filename(filename: str) -> str:
    if not filename:
        return ""
    import re
    base = re.sub(r'\.[a-zA-Z0-9]+$', '', filename)
    base = re.sub(r'^\d+[\s_-]*', '', base)
    base = re.sub(r'[\s_-]*(?:resume|cv|curriculum_vitae|profile)[\s_-]*', '', base, flags=re.IGNORECASE)
    name = re.sub(r'[_-]', ' ', base).strip()
    words = name.split()
    if 1 <= len(words) <= 5 and all(w.isalpha() for w in words):
        return ' '.join(w.capitalize() for w in words)
    return ""

def _extract_profile_rule_based(raw_text: str, default_email: str, filename: str = "") -> dict:
    import re
    lines = [l.strip() for l in raw_text.splitlines() if l.strip()]
    
    # 1. Email extraction
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', raw_text)
    email = email_match.group(0) if email_match else default_email
    
    # 2. Phone extraction
    phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}', raw_text)
    phone = phone_match.group(0) if phone_match else ""
    
    # 3. Name & Title extraction from filename and top text lines
    full_name = _extract_name_from_filename(filename)
    title = ""
    for line in lines[:10]:
        clean = line.strip()
        if "@" in clean or any(char.isdigit() for char in clean):
            continue
        words = clean.split()
        if 2 <= len(words) <= 4 and all(w.isalpha() for w in words):
            if not any(w.lower() in ["resume", "curriculum", "vitae", "contact", "email", "phone", "profile", "summary", "experience", "education", "skills"] for w in words):
                if not full_name or full_name == "Candidate":
                    full_name = clean
                    continue
        if full_name and not title and len(clean) < 60:
            if not any(w.lower() in ["contact", "email", "phone", "profile", "summary", "experience", "education", "skills"] for w in clean.split()):
                title = clean
                
    if not full_name:
        full_name = "Candidate"
                
    # 4. Skills extraction
    found_skills = []
    lower_text = raw_text.lower()
    for skill in COMMON_SKILLS:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, lower_text):
            found_skills.append({
                "id": f"skl_{uuid.uuid4().hex[:8]}",
                "name": skill,
                "category": "hard"
            })
            
    # 5. Experience extraction using date patterns
    experiences = []
    date_pattern = re.compile(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4})\s*(?:-|–|to)\s*((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|\d{4}|Present|Current)', re.IGNORECASE)
    exp_matches = list(date_pattern.finditer(raw_text))
    for i, match in enumerate(exp_matches[:5]):
        start_pos = match.start()
        preceding = raw_text[:start_pos].strip().splitlines()
        comp_title = preceding[-1].strip() if preceding else "Company / Role"
        parts = re.split(r'[-–|,@]', comp_title)
        role = parts[0].strip() if len(parts) > 0 else "Role"
        company = parts[1].strip() if len(parts) > 1 else comp_title
        
        end_pos = exp_matches[i + 1].start() if i + 1 < len(exp_matches) else start_pos + 600
        section_chunk = raw_text[match.end():end_pos]
        bullet_lines = [b.strip().lstrip("•-*– ") for b in section_chunk.splitlines() if len(b.strip()) > 15]
        
        bullets = [{"id": f"blt_{uuid.uuid4().hex[:8]}", "text": b} for b in bullet_lines[:4]]
        if not bullets:
            bullets = [{"id": f"blt_{uuid.uuid4().hex[:8]}", "text": "Responsible for core engineering and key product deliverables."}]
            
        experiences.append({
            "id": f"exp_{uuid.uuid4().hex[:8]}",
            "company": company,
            "title": role,
            "start_date": match.group(1),
            "end_date": match.group(2),
            "bullets": bullets
        })
        
    # 6. Education extraction
    education = []
    edu_keywords = ["bachelor", "master", "b.s", "b.a", "b.sc", "m.s", "m.sc", "ph.d", "degree", "university", "college", "institute", "polytechnic"]
    for line in lines:
        lower_line = line.lower()
        if any(kw in lower_line for kw in edu_keywords) and len(line) < 120:
            if not any(header == lower_line for header in ["education", "academic background"]):
                year_match = re.search(r'\b(20\d{2}|19\d{2})\b', line)
                year_str = year_match.group(0) if year_match else ""
                
                parts = re.split(r'[-–|,]', line)
                deg = parts[0].strip() if len(parts) > 0 else line
                inst = parts[1].strip() if len(parts) > 1 else line
                
                education.append({
                    "id": f"edu_{uuid.uuid4().hex[:8]}",
                    "institution": inst,
                    "degree": deg,
                    "field_of_study": "",
                    "start_year": "",
                    "end_year": year_str
                })
                if len(education) >= 3:
                    break
                
    return {
        "full_name": full_name,
        "title": title or "Software Professional",
        "email": email,
        "phone": phone,
        "location": "",
        "bio": f"{title or 'Professional'} with demonstrated technical and engineering experience.",
        "experiences": experiences,
        "education": education,
        "skills": found_skills
    }

class ProfileService:
    @staticmethod
    async def get_or_create_profile(db, user_id: str) -> MasterProfile:
        result = await db.execute(select(MasterProfile).where(MasterProfile.user_email == user_id))
        profile = result.scalars().first()
        if not profile:
            from app.models.user import User as UserModel
            user_res = await db.execute(select(UserModel).where(UserModel.email == user_id))
            user_obj = user_res.scalars().first()
            default_name = (user_obj.full_name if user_obj and user_obj.full_name else "Candidate")

            profile = MasterProfile(
                user_email=user_id,
                full_name=default_name,
                title="",
                email=user_id,
                phone="",
                location="",
                bio="",
                linkedin_url="",
                github_url="",
                portfolio_url="",
                experiences=[],
                education=[],
                skills=[]
            )
            db.add(profile)
            await db.commit()
            await db.refresh(profile)
        return profile

    @staticmethod
    async def get_full_profile(db, user_id: str):
        result = await db.execute(select(MasterProfile).where(MasterProfile.user_email == user_id))
        profile = result.scalars().first()
        if not profile:
            return {
                "id": user_id,
                "full_name": "",
                "title": "",
                "email": user_id,
                "phone": "",
                "location": "",
                "bio": "",
                "linkedin_url": "",
                "github_url": "",
                "portfolio_url": "",
                "experiences": [],
                "education": [],
                "skills": []
            }
        
        return {
            "id": profile.id,
            "full_name": profile.full_name or "",
            "title": profile.title or "",
            "email": profile.email or user_id,
            "phone": profile.phone or "",
            "location": profile.location or "",
            "bio": profile.bio or "",
            "linkedin_url": profile.linkedin_url or "",
            "github_url": profile.github_url or "",
            "portfolio_url": profile.portfolio_url or "",
            "experiences": profile.experiences or [],
            "education": profile.education or [],
            "skills": profile.skills or [],
        }

    @staticmethod
    async def update_details(db, user_id: str, data):
        profile = await ProfileService.get_or_create_profile(db, user_id)
        if data.full_name is not None:
            profile.full_name = data.full_name
        if data.title is not None:
            profile.title = data.title
        if data.email is not None:
            profile.email = data.email
        if data.phone is not None:
            profile.phone = data.phone
        if data.location is not None:
            profile.location = data.location
        if data.bio is not None:
            profile.bio = data.bio
        await db.commit()
        await db.refresh(profile)
        return await ProfileService.get_full_profile(db, user_id)

    @staticmethod
    async def update_links(db, user_id: str, data):
        profile = await ProfileService.get_or_create_profile(db, user_id)
        if data.linkedin_url is not None:
            profile.linkedin_url = str(data.linkedin_url)
        if data.github_url is not None:
            profile.github_url = str(data.github_url)
        if data.portfolio_url is not None:
            profile.portfolio_url = str(data.portfolio_url)
        await db.commit()
        await db.refresh(profile)
        return await ProfileService.get_full_profile(db, user_id)

    @staticmethod
    async def create_experience(db, user_id: str, data):
        profile = await ProfileService.get_or_create_profile(db, user_id)
        current_exps = list(profile.experiences or [])
        new_exp = data.model_dump() if hasattr(data, "model_dump") else data.dict()
        if not new_exp.get("id"):
            new_exp["id"] = f"exp_{uuid.uuid4().hex[:8]}"
        current_exps.append(new_exp)
        profile.experiences = current_exps
        await db.commit()
        await db.refresh(profile)
        return await ProfileService.get_full_profile(db, user_id)

    @staticmethod
    async def update_experience(db, user_id: str, exp_id: str, data):
        profile = await ProfileService.get_or_create_profile(db, user_id)
        current_exps = list(profile.experiences or [])
        updated_exps = []
        updated_dict = data.model_dump() if hasattr(data, "model_dump") else data.dict()
        for e in current_exps:
            if e.get("id") == exp_id:
                updated_dict["id"] = exp_id
                updated_exps.append(updated_dict)
            else:
                updated_exps.append(e)
        profile.experiences = updated_exps
        await db.commit()
        await db.refresh(profile)
        return await ProfileService.get_full_profile(db, user_id)

    @staticmethod
    async def delete_experience(db, user_id: str, exp_id: str):
        profile = await ProfileService.get_or_create_profile(db, user_id)
        current_exps = list(profile.experiences or [])
        profile.experiences = [e for e in current_exps if e.get("id") != exp_id]
        await db.commit()
        await db.refresh(profile)
        return await ProfileService.get_full_profile(db, user_id)

    @staticmethod
    async def update_skills(db, user_id: str, data):
        profile = await ProfileService.get_or_create_profile(db, user_id)
        skills_list = []
        for s in data:
            s_dict = s.model_dump() if hasattr(s, "model_dump") else s.dict() if hasattr(s, "dict") else s
            skills_list.append(s_dict)
        profile.skills = skills_list
        await db.commit()
        await db.refresh(profile)
        return await ProfileService.get_full_profile(db, user_id)

    @staticmethod
    async def create_education(db, user_id: str, data):
        profile = await ProfileService.get_or_create_profile(db, user_id)
        current_edu = list(profile.education or [])
        new_edu = data.model_dump() if hasattr(data, "model_dump") else data.dict()
        if not new_edu.get("id"):
            new_edu["id"] = f"edu_{uuid.uuid4().hex[:8]}"
        current_edu.append(new_edu)
        profile.education = current_edu
        await db.commit()
        await db.refresh(profile)
        return await ProfileService.get_full_profile(db, user_id)

    @staticmethod
    async def update_education(db, user_id: str, edu_id: str, data):
        profile = await ProfileService.get_or_create_profile(db, user_id)
        current_edu = list(profile.education or [])
        updated_edu = []
        updated_dict = data.model_dump() if hasattr(data, "model_dump") else data.dict()
        for ed in current_edu:
            if ed.get("id") == edu_id:
                updated_dict["id"] = edu_id
                updated_edu.append(updated_dict)
            else:
                updated_edu.append(ed)
        profile.education = updated_edu
        await db.commit()
        await db.refresh(profile)
        return await ProfileService.get_full_profile(db, user_id)

    @staticmethod
    async def delete_education(db, user_id: str, edu_id: str):
        profile = await ProfileService.get_or_create_profile(db, user_id)
        current_edu = list(profile.education or [])
        profile.education = [ed for ed in current_edu if ed.get("id") != edu_id]
        await db.commit()
        await db.refresh(profile)
        return await ProfileService.get_full_profile(db, user_id)

    @staticmethod
    async def parse_cv(file_bytes: bytes, filename: str, current_user: str, db) -> dict:
        raw_text = ""
        
        # 1. Extract text based on file format
        safe_filename = (filename or "").lower()
        if safe_filename.endswith('.pdf'):
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                for page in doc:
                    raw_text += page.get_text("text") + "\n"
            except Exception as e:
                print(f"PyMuPDF error: {e}")
        elif safe_filename.endswith('.docx'):
            try:
                doc = Document(io.BytesIO(file_bytes))
                # 1. Header paragraphs from all sections
                for section in doc.sections:
                    if section.header:
                        for p in section.header.paragraphs:
                            if p.text.strip():
                                raw_text += p.text + "\n"
                # 2. Main body paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        raw_text += para.text + "\n"
                # 3. Table cells (crucial for modern 2-column resume templates)
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = []
                        for cell in row.cells:
                            cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                            if cell_text and cell_text not in row_texts:
                                row_texts.append(cell_text)
                        if row_texts:
                            raw_text += " | ".join(row_texts) + "\n"
            except Exception as e:
                print(f"python-docx error: {e}")
        else:
            raw_text = file_bytes.decode('utf-8', errors='ignore')
            
        if not raw_text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Could not extract any readable text from the uploaded document."
            )

        # 2. Call Gemini asynchronously to parse and structure the text
        prompt = f"""
        You are an expert ATS resume parser.
        Your task is to accurately extract ALL profile information from the raw resume text into the required JSON schema.
        
        CRITICAL EXTRACTION REQUIREMENTS:
        1. "full_name": Extract the candidate's actual personal name from the resume header (usually at the very top or in the filename header). Do NOT return generic placeholders like "Candidate" or "User".
        2. "title": Extract the candidate's current or target job title/headline (e.g. "Full-Stack Engineer", "Software Architect").
        3. "email", "phone", "location": Extract all contact information.
        4. "bio": Extract or construct a concise 1-2 sentence professional summary of the candidate.
        5. "education": Extract all degree programs, universities/colleges, fields of study, and graduation years. Be thorough and capture every educational qualification.
        6. "experiences": Extract all work experience entries with company names, job titles, start and end dates, and individual bullet point accomplishments.
        7. "skills": Extract all technical skills, frameworks, tools, and languages mentioned.
        
        RAW RESUME TEXT:
        {raw_text[:30000]}
        """
        
        try:
            response_text = await asyncio.to_thread(_call_gemini_extraction_sync, prompt)
            if not response_text.strip():
                raise ValueError("Empty response from extraction model")
            result_data = json.loads(response_text)
        except Exception as e:
            print(f"[Parser Fallback] AI parser unavailable ({e}), using instant local rule-based extractor...")
            result_data = _extract_profile_rule_based(raw_text, current_user, filename=filename)
        
        # 3. Save to database safely without erasing existing data if parse is empty
        profile = await ProfileService.get_or_create_profile(db, current_user)
        extracted_name = result_data.get("full_name") or ""
        if not extracted_name or extracted_name == "Candidate":
            from_file = _extract_name_from_filename(filename)
            if from_file:
                extracted_name = from_file
        if extracted_name and extracted_name != "Candidate":
            profile.full_name = extracted_name

        if result_data.get("title") and result_data["title"].strip():
            profile.title = result_data["title"]
        if result_data.get("email") and result_data["email"].strip():
            profile.email = result_data["email"]
        if result_data.get("phone") and result_data["phone"].strip():
            profile.phone = result_data["phone"]
        if result_data.get("location") and result_data["location"].strip():
            profile.location = result_data["location"]
        if result_data.get("bio") and result_data["bio"].strip():
            profile.bio = result_data["bio"]
        
        if result_data.get("experiences"):
            profile.experiences = result_data["experiences"]
        if result_data.get("education"):
            profile.education = result_data["education"]
        if result_data.get("skills"):
            profile.skills = result_data["skills"]
            
        await db.commit()
        await db.refresh(profile)
        
        return await ProfileService.get_full_profile(db, current_user)

    @staticmethod
    async def parse_linkedin(url: str, current_user: str, db) -> dict:
        import httpx
        
        # SSRF and URL validation
        parsed_url = urlparse(url)
        if parsed_url.scheme not in ("http", "https") or "linkedin.com" not in parsed_url.netloc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid URL. Please provide a valid public LinkedIn profile URL."
            )
        
        # Guard against private/internal IPs
        hostname = parsed_url.hostname or ""
        if hostname in ("localhost", "127.0.0.1", "0.0.0.0", "169.254.169.254") or hostname.startswith("192.168.") or hostname.startswith("10."):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Restricted destination address."
            )

        html_content = ""
        try:
            async with httpx.AsyncClient() as client_http:
                headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
                response = await client_http.get(url, headers=headers, follow_redirects=True, timeout=12.0)
                html_content = response.text
        except Exception as e:
            print(f"Failed to fetch LinkedIn URL: {e}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Could not reach the LinkedIn profile URL. Ensure the profile is public."
            )
            
        prompt = f'''
        You are an expert ATS parser. Extract the full profile from this raw LinkedIn HTML.
        LinkedIn heavily obfuscates their HTML, so look for JSON-LD scripts or basic text blocks.
        Structure the candidate name, title, experiences, education, and skills. Be highly accurate.
        If you hit an Authwall or cannot find data, return empty structures.
        
        RAW HTML:
        {html_content[:30000]}
        '''
        
        try:
            response_text = await asyncio.to_thread(_call_gemini_extraction_sync, prompt)
            if not response_text.strip():
                raise ValueError("Empty response from extraction model")
            result_data = json.loads(response_text)
        except Exception as e:
            print(f"[LinkedIn Fallback] AI parser unavailable ({e}), using instant local rule-based extractor...")
            result_data = _extract_profile_rule_based(html_content, current_user)
        
        # Save to database safely without erasing existing data if parse is empty
        profile = await ProfileService.get_or_create_profile(db, current_user)
        if result_data.get("full_name") and result_data["full_name"].strip() and result_data["full_name"] != "Candidate":
            profile.full_name = result_data["full_name"]
        if result_data.get("title") and result_data["title"].strip():
            profile.title = result_data["title"]
        if result_data.get("email") and result_data["email"].strip():
            profile.email = result_data["email"]
        if result_data.get("phone") and result_data["phone"].strip():
            profile.phone = result_data["phone"]
        if result_data.get("location") and result_data["location"].strip():
            profile.location = result_data["location"]
        if result_data.get("bio") and result_data["bio"].strip():
            profile.bio = result_data["bio"]
        profile.linkedin_url = url
        
        if result_data.get("experiences"):
            profile.experiences = result_data["experiences"]
        if result_data.get("education"):
            profile.education = result_data["education"]
        if result_data.get("skills"):
            profile.skills = result_data["skills"]
            
        await db.commit()
        await db.refresh(profile)
        
        return await ProfileService.get_full_profile(db, current_user)
